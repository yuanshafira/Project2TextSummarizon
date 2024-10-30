from flask import Flask, request, jsonify, send_from_directory
import requests
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from urllib.parse import urlparse
import os
import docx
import base64

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads_model/.media/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Limit to 16MB

if not os.path.exists(UPLOAD_FOLDER):
    print("tidak ada folder .media")

tokenizer = AutoTokenizer.from_pretrained("rowjak/bert-indonesian-news-summarization")
model = AutoModelForSeq2SeqLM.from_pretrained("rowjak/bert-indonesian-news-summarization")

# Move model to GPU if available
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

def extract_text_from_docx(file_path):
    """Extract text from a .docx file."""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def save_text_to_docx(text, output_path):
    """Simpan teks ke file .docx baru."""
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(output_path)


@app.route('/uploads/<path:filename>')
def download_file(filename):
    return send_from_directory('uploads_model/uploads', filename)

@app.route('/webhook', methods=['POST'])
def webhook():
    # Mengambil data JSON dari request
    data = request.get_json()

    # Memeriksa apakah payload ada dalam data
    if 'payload' not in data:
        return jsonify({'error': 'Payload not found'}), 400


    if data['payload'].get('hasMedia', None) == True:

        media_data = data['payload'].get('media', {})

        filenames = media_data.get('filename', None)
        url = media_data.get('url', None)


        parsed_url = urlparse(url).path

        filename = os.path.basename(parsed_url)

        file_path = f'.media/{filename}'
        extracted_text = extract_text_from_docx(file_path)

        input_ids = tokenizer.encode(extracted_text, return_tensors='pt').to(device)
        summary_ids = model.generate(input_ids,
                                     min_length=20,
                                     max_length=80,
                                     num_beams=10,
                                     repetition_penalty=2.5,
                                     length_penalty=1.0,
                                     early_stopping=True,
                                     no_repeat_ngram_size=2,
                                     use_cache=True,
                                     do_sample=True,
                                     temperature=0.8,
                                     top_k=50,
                                     top_p=0.95)
        summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

        filenames = filenames.replace(" ", "-")

        save_text_to_docx(summary_text,f"uploads/summary-{filenames}")

        from_number = data['payload'].get('from', None)

        response = requests.post(
            'https://xxxx/api/sendFile',
            headers={
                'Content-Type': 'application/json; charset=utf-8',
                'Accept': 'application/json',
                'X-Api-Key': 'xxxx'
            },
            json={
                'chatId': from_number,
                'caption' : '',
                'session': 'NoamChomsky',
                'file' : {
                    'mimetype' : 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'url' : f"{request.host_url}/summary-{filenames}",
                    'filename' : 'user.docx'
                }
            },
            verify=False
        )

        return jsonify(response.json()), response.status_code

    else:
        body = data['payload'].get('body', None)
        input_ids = tokenizer.encode(body, return_tensors='pt').to(device)
        summary_ids = model.generate(input_ids,
                                     min_length=20,
                                     max_length=80,
                                     num_beams=10,
                                     repetition_penalty=2.5,
                                     length_penalty=1.0,
                                     early_stopping=True,
                                     no_repeat_ngram_size=2,
                                     use_cache=True,
                                     do_sample=True,
                                     temperature=0.8,
                                     top_k=50,
                                     top_p=0.95)
        summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

        from_number = data['payload'].get('from', None)

        response = requests.post(
            'https://xxxx/api/sendText',
            headers={
                'Content-Type': 'application/json; charset=utf-8',
                'Accept': 'application/json',
                'X-Api-Key': 'xxxx'
            },
            json={  # Menggunakan parameter `json` untuk mengirim raw JSON
                'chatId': from_number,
                'text': summary_text,
                'session': 'NoamChomsky'
            },
            verify=False  # Ini sesuai dengan withoutVerifying() di PHP
        )
        # Mengembalikan respons dari permintaan HTTP
        return jsonify(response.json()), response.status_code


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)

